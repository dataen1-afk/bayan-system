"""
Grant Agreement Generator
Uses the official BAYAN template to generate professional Grant Agreement PDFs
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor, black, white
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT, TA_JUSTIFY
from pathlib import Path
import os
import copy
import arabic_reshaper
from bidi.algorithm import get_display

ROOT_DIR = Path(__file__).parent
ASSETS_DIR = ROOT_DIR / "assets"
FONTS_DIR = ROOT_DIR / "fonts"
TEMPLATE_PATH = ASSETS_DIR / "grant_agreement_template.docx"

# Register Arabic fonts
try:
    pdfmetrics.registerFont(TTFont('Amiri', str(FONTS_DIR / 'Amiri-Regular.ttf')))
    pdfmetrics.registerFont(TTFont('Amiri-Bold', str(FONTS_DIR / 'Amiri-Bold.ttf')))
    ARABIC_FONT_REGISTERED = True
except:
    ARABIC_FONT_REGISTERED = False

# Brand colors
NAVY_BLUE = HexColor('#1e3a5f')
GOLD = HexColor('#c9a227')


def process_arabic(text):
    """Process Arabic text for proper display"""
    if not text:
        return text
    try:
        reshaped = arabic_reshaper.reshape(str(text))
        return get_display(reshaped)
    except:
        return str(text)


def generate_grant_agreement_pdf(agreement_data: dict, output_path: str) -> str:
    """
    Generate a Grant Agreement PDF using the official BAYAN template
    
    Args:
        agreement_data: Dictionary containing agreement information
        output_path: Path to save the PDF
    
    Returns:
        Path to generated PDF
    """
    
    # Create PDF using ReportLab with proper formatting
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Header style
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=6
    )
    
    # Title style
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=12,
        spaceBefore=12,
        textColor=NAVY_BLUE,
        fontName='Helvetica-Bold'
    )
    
    # Arabic title style
    arabic_title_style = ParagraphStyle(
        'ArabicTitle',
        parent=styles['Heading1'],
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=20,
        textColor=NAVY_BLUE,
        fontName='Amiri-Bold' if ARABIC_FONT_REGISTERED else 'Helvetica-Bold'
    )
    
    # Section header style
    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading2'],
        fontSize=11,
        spaceBefore=12,
        spaceAfter=6,
        textColor=NAVY_BLUE,
        fontName='Helvetica-Bold'
    )
    
    # Body style
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
        leading=14
    )
    
    # Bullet style
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        leftIndent=20,
        bulletIndent=10
    )
    
    # Build content
    story = []
    
    # Header with logo
    logo_path = ASSETS_DIR / "bayan-logo.png"
    if logo_path.exists():
        try:
            story.append(RLImage(str(logo_path), width=2*cm, height=1.5*cm))
        except:
            pass
    
    # Company header
    story.append(Paragraph("BAYAN AUDITING & CONFORMITY", header_style))
    story.append(Paragraph("بيان للتحقق والمطابقة", ParagraphStyle('ArabicHeader', parent=header_style, fontName='Amiri' if ARABIC_FONT_REGISTERED else 'Helvetica')))
    story.append(Paragraph("3879 Al Khadar Street, Riyadh, 12282, Saudi Arabia", header_style))
    story.append(Spacer(1, 0.5*cm))
    
    # Main Title
    story.append(Paragraph("GRANT AGREEMENT", title_style))
    ar_title = process_arabic("اتفاقية المنح") if ARABIC_FONT_REGISTERED else "اتفاقية المنح"
    story.append(Paragraph(ar_title, arabic_title_style))
    story.append(Spacer(1, 0.5*cm))
    
    # Extract data
    org_name = agreement_data.get('organization_name', agreement_data.get('organizationName', ''))
    org_address = agreement_data.get('organization_address', agreement_data.get('organizationAddress', ''))
    standards = agreement_data.get('standards', [])
    scope = agreement_data.get('scope', '')
    contact_name = agreement_data.get('contact_name', agreement_data.get('contactName', ''))
    
    # Section 1: Parties
    story.append(Paragraph("1. Parties to the Agreement", section_style))
    story.append(Paragraph(
        "This Grant Agreement (\"Agreement\") is made between:",
        body_style
    ))
    story.append(Paragraph(
        "<b>• Certification Body:</b> BAYAN AUDITING & CONFORMITY (BAC) Arabia Limited Certification Body (hereafter referred to as \"the Certification Body\" or \"CB\")",
        bullet_style
    ))
    story.append(Paragraph(
        "Address: 3879 Al Khadar Street, Riyadh, 12282, Saudi Arabia",
        ParagraphStyle('Address', parent=bullet_style, leftIndent=40)
    ))
    story.append(Paragraph(
        f"<b>• Client Facility:</b> {org_name} (hereafter referred to as \"the Client\")",
        bullet_style
    ))
    story.append(Paragraph(
        f"Address: {org_address}",
        ParagraphStyle('Address', parent=bullet_style, leftIndent=40)
    ))
    
    # Section 2: Purpose and Scope
    story.append(Paragraph("2. Purpose and Scope", section_style))
    story.append(Paragraph(
        "This Agreement defines the legally enforceable terms and conditions under which the CB provides certification services for the following management system standard(s):",
        body_style
    ))
    
    # Standards checkboxes
    all_standards = ['ISO 9001', 'ISO 14001', 'ISO 45001', 'ISO 22000', 'ISO 22301', 'ISO 27001']
    standards_text = ""
    for std in all_standards:
        checked = "☑" if std in standards else "☐"
        standards_text += f"{checked} {std}   "
    story.append(Paragraph(standards_text, body_style))
    
    if scope:
        story.append(Paragraph(f"<b>Scope:</b> {scope}", body_style))
    
    # Section 3: Legal Enforceability
    story.append(Paragraph("3. Legal Enforceability", section_style))
    story.append(Paragraph("• This Agreement is legally binding.", bullet_style))
    story.append(Paragraph("• The CB retains full authority to grant, maintain, reduce, suspend, or withdraw certification.", bullet_style))
    story.append(Paragraph("• Certification activities shall comply with ISO/IEC 17021-1:2015 and applicable IAF Mandatory Documents.", bullet_style))
    
    # Section 4: Certification Body Responsibilities
    story.append(Paragraph("4. Certification Body Responsibilities", section_style))
    story.append(Paragraph("The Certification Body agrees to:", body_style))
    story.append(Paragraph("• Maintain impartiality and avoid conflicts of interest", bullet_style))
    story.append(Paragraph("• Assign competent auditors for the applicable scheme and sector", bullet_style))
    story.append(Paragraph("• Notify clients of changes in certification requirements", bullet_style))
    story.append(Paragraph("• Respect confidentiality of client data", bullet_style))
    story.append(Paragraph("• Handle appeals and complaints transparently", bullet_style))
    
    # Section 5: Client Responsibilities
    story.append(Paragraph("5. Client Responsibilities", section_style))
    story.append(Paragraph("The Client agrees to:", body_style))
    story.append(Paragraph("• Comply with applicable certification requirements", bullet_style))
    story.append(Paragraph("• Provide unrestricted access to premises, personnel, documentation for audit purposes", bullet_style))
    story.append(Paragraph("• Maintain system performance and notify CB of changes affecting certification", bullet_style))
    story.append(Paragraph("• Use certification status in accordance with CB requirements", bullet_style))
    story.append(Paragraph("• Cease all use of certificates upon suspension or withdrawal", bullet_style))
    
    # Section 6: Use of Certification
    story.append(Paragraph("6. Use and Misuse of Certification", section_style))
    story.append(Paragraph("The Client shall:", body_style))
    story.append(Paragraph("a) Follow CB guidance when referencing certification", bullet_style))
    story.append(Paragraph("b) Avoid any misleading statements regarding certification", bullet_style))
    story.append(Paragraph("c) Not misuse certification documents or marks", bullet_style))
    story.append(Paragraph("d) Discontinue use upon withdrawal or suspension", bullet_style))
    
    # Section 7: Confidentiality
    story.append(Paragraph("7. Confidentiality", section_style))
    story.append(Paragraph(
        "The CB shall manage all information obtained during certification as confidential, except where disclosure is required by law or approved by the client.",
        body_style
    ))
    
    # Section 8-14: Brief summaries
    story.append(Paragraph("8. Information Provided by the CB", section_style))
    story.append(Paragraph("The CB shall provide the client with certification process details, normative requirements, applicable fees, and complaints/appeals procedures.", body_style))
    
    story.append(Paragraph("9. Fees and Payment", section_style))
    story.append(Paragraph("All fees are as per the official proposal. Invoices are payable within 14 days. Certification is issued only after full payment.", body_style))
    
    story.append(Paragraph("10. Notification of Changes", section_style))
    story.append(Paragraph("The Client shall inform the CB of any changes affecting certification without delay.", body_style))
    
    story.append(Paragraph("11. Suspension, Withdrawal, and Termination", section_style))
    story.append(Paragraph("Either party may terminate with two months' written notice. Upon termination, the Client shall cease all use of certification.", body_style))
    
    story.append(Paragraph("12. Liability and Indemnity", section_style))
    story.append(Paragraph("Maximum liability is limited to the total fee paid for certification.", body_style))
    
    story.append(Paragraph("13. Force Majeure", section_style))
    story.append(Paragraph("Neither party is liable for non-performance due to events beyond their control.", body_style))
    
    story.append(Paragraph("14. Governing Law", section_style))
    story.append(Paragraph("This Agreement is governed by the laws of the Kingdom of Saudi Arabia.", body_style))
    
    # Acknowledgments
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph("Acknowledgments:", section_style))
    story.append(Paragraph("☑ We have reviewed and understood the BAYAN certification process.", bullet_style))
    story.append(Paragraph("☑ We agree to the rules of certification mark usage.", bullet_style))
    story.append(Paragraph("☑ We will notify BAYAN of any significant changes to the management system.", bullet_style))
    
    # Signatures section
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("Acceptance and Signatures", section_style))
    
    # Signature table
    sig_data = [
        ["For BAYAN AUDITING & CONFORMITY", "For Client Facility"],
        ["Name: Islam Abd El-Aal", f"Name: {contact_name}"],
        ["Position: Director", "Position: ________________"],
        ["Signature: ________________", "Signature: ________________"],
        ["Date: ________________", "Date: ________________"]
    ]
    
    sig_table = Table(sig_data, colWidths=[8*cm, 8*cm])
    sig_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.append(sig_table)
    
    # Build PDF
    doc.build(story)
    
    return output_path


def fill_docx_template(agreement_data: dict, output_path: str) -> str:
    """
    Fill the DOCX template with agreement data
    
    Args:
        agreement_data: Dictionary containing agreement information
        output_path: Path to save the filled document
    
    Returns:
        Path to generated DOCX
    """
    
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    
    # Load template
    doc = Document(str(TEMPLATE_PATH))
    
    # Extract data
    org_name = agreement_data.get('organization_name', agreement_data.get('organizationName', ''))
    org_address = agreement_data.get('organization_address', agreement_data.get('organizationAddress', ''))
    standards = agreement_data.get('standards', [])
    scope = agreement_data.get('scope', '')
    contact_name = agreement_data.get('contact_name', agreement_data.get('contactName', ''))
    
    # Replace placeholders in the document
    for paragraph in doc.paragraphs:
        if 'Client Organization' in paragraph.text:
            paragraph.text = paragraph.text.replace('Client Organization', org_name)
        if 'XXXXXXXXXXXXXXXXXXX' in paragraph.text:
            paragraph.text = paragraph.text.replace('XXXXXXXXXXXXXXXXXXX', org_address)
        if 'xxxxx' in paragraph.text.lower():
            paragraph.text = paragraph.text.replace('xxxxx', org_name)
            paragraph.text = paragraph.text.replace('XXXXX', org_name)
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if 'Client Organization' in paragraph.text:
                        paragraph.text = paragraph.text.replace('Client Organization', org_name)
                    if 'XXXXXXXXXXXXXXXXXXX' in paragraph.text:
                        paragraph.text = paragraph.text.replace('XXXXXXXXXXXXXXXXXXX', org_address)
    
    # Save filled document
    doc.save(output_path)
    
    return output_path
